# core/utils.py
import random
def rank_candidates(job_description, cvs):
    """
    Rank candidates based on the provided job description and CVs.
    """

    ranked_results = []
    for cv in cvs:
        # التأكد من أن المهارات ليست 0
        matching_skills = len(set(cv.skills).intersection(job_description.description.split()))
        skills_score = (matching_skills * 0.4) if matching_skills > 0 else 0.1

        # التأكد من أن سنوات الخبرة محسوبة
        experience_score = (cv.total_experience_years / 10 * 0.3) if cv.total_experience_years > 0 else 0.1

        # التأكد من حساب التعليم بشكل صحيح
        education_score = 0.3 if cv.degrees else 0

        # حساب النتيجة النهائية
        overall_score = skills_score + experience_score + education_score

        overall_score_random = random.random()

        # التأكد من عدم إعادة overall_score كـ 0 أبدًا
        if overall_score == 0:
            overall_score = overall_score_random  # قيمة احتياطية صغيرة جدًا
            print(overall_score)
        elif overall_score <=0.5:
            overall_score = overall_score + 0.35

        recommendation = "Let's have interview" if overall_score > 0.65 else "save for future"

        ranked_results.append({
            "cv_id":cv.id,
            "name": cv.name,
            "email": cv.email,
            "gender": cv.gender,
            "phone_number":cv.phone_number,
            "nationality":cv.nationality,
            "linkedin":cv.linkedin,
            "degrees":cv.degrees,
            "designation":cv.designation,
            "last_company":cv.last_company,
            "skills":cv.skills,
            "total_experience_years":cv.total_experience_years,
            "overall_score": overall_score,
            "recommendation": recommendation
        })

    # ترتيب المرشحين حسب الدرجة النهائية تنازليًا
    ranked_results.sort(key=lambda x: x['overall_score'], reverse=True)

    return ranked_results

import os
import torch
import transformers
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from docx import Document
import pdfplumber

# تحميل نموذج LLAMA
model_id = "meta-llama/Meta-Llama-3.1-8B-Instruct"
os.environ["TOKENIZERS_PARALLELISM"] = "false"

pipeline = transformers.pipeline(
    "text-generation",
    model=model_id,
    model_kwargs={"torch_dtype": torch.bfloat16},
    device_map="auto",
)

def extract_text_from_file(file):
    """ استخراج النصوص من ملفات السير الذاتية """
    try:
        if file.name.endswith('.docx'):
            document = Document(file)
            return '\n'.join([paragraph.text for paragraph in document.paragraphs])
        elif file.name.endswith('.pdf'):
            with pdfplumber.open(file) as pdf:
                return '\n'.join(page.extract_text() for page in pdf.pages if page.extract_text())
        else:
            return file.read().decode('utf-8')
    except Exception as e:
        return f"Error reading file: {str(e)}"

@csrf_exempt
def extract_resume_features(request):
    """ استخراج مميزات السيرة الذاتية """
    if request.method == 'POST' and request.FILES.get('resume'):
        file = request.FILES['resume']
        text = extract_text_from_file(file)

        prompts = {
            "Name": "Provide the name of the person.",
            "Email": "Provide the email address of the person.",
            "Phone_Number": "Provide the phone number of the person.",
            "LinkedIn": "Provide the LinkedIn profile of the person.",
            "Skills": "List all the tangible skills from the resume.",
            "Years_of_Experience": "Calculate total years of professional experience."
        }

        extracted_data = {}

        for key, prompt in prompts.items():
            messages = [
                {"role": "system", "content": "Provide direct answers without explanation."},
                {"role": "user", "content": text},
                {"role": "user", "content": prompt}
            ]

            output = pipeline(messages, max_new_tokens=100, temperature=0.5, top_k=50, top_p=0.9)
            extracted_data[key] = output[0]["generated_text"][-1]['content']

        return JsonResponse({"resume_data": extracted_data}, status=200)

    return JsonResponse({"error": "Invalid request"}, status=400)

import re
import pdfplumber
from docx import Document


def extract_cv_data(file):
    data = {
        "name": "",
        "email": "",
        "gender": "",
        "phone_number": "",
        "nationality": "",
        "linkedin": "",
        "degrees": "",
        "college_name": "",
        "designation": "",
        "last_company": "",
        "skills": [],
        "total_experience_years": 0,
    }

    try:
        # استخراج النص من الملف
        if file.name.endswith('.pdf'):
            with pdfplumber.open(file) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        elif file.name.endswith(('.doc', '.docx')):
            doc = Document(file)
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            text = file.read().decode('utf-8')

        # **تحسين الأنماط**
        name_pattern = r"Name:\s*([\w\s-]+)"
        email_pattern = r"[\w\.-]+@[\w\.-]+\.\w+"
        phone_pattern = r"(?:\+?\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}"
        gender_pattern = r"Gender:\s*(\w+)"
        nationality_pattern = r"Nationality:\s*([\w\s]+)"
        linkedin_pattern = r"(linkedin\.com/in/[\w-]+)"
        degrees_pattern = r"Bachelor of [\w\s]+ in ([\w\s]+)"
        college_name_pattern = r"(?:University|Institute|College|School of [\w\s]+)"

        # تحسين البحث عن الوظيفة والشركة الأخيرة
        designation_pattern = r"(?:Position|Designation|Role):\s*([\w\s&-]+)"
        last_company_pattern = r"(?:Company|Employer|Organization):\s*([\w\s&-]+)"

        # البحث في قسم "Work Experience"
        job_experience_pattern = r"([\w\s&-]+)\s+(?:at|في|with)\s+([\w\s&-]+(?:\s+[\w\s&-]+)?)"

        skills_pattern = r"(?:Skills|Technical Skills|Key Skills)\s*[:\n]([\s\S]*?)\n(?:Experience|Education|Projects|Work Experience|Certifications)"
        experience_pattern = r"(\d{4})\s*-\s*(\d{4}|Present|Now)"

        # **استخراج البيانات**
        name_match = re.search(name_pattern, text, re.IGNORECASE)
        email_match = re.search(email_pattern, text, re.IGNORECASE)
        phone_match = re.search(phone_pattern, text, re.IGNORECASE)
        gender_match = re.search(gender_pattern, text, re.IGNORECASE)
        nationality_match = re.search(nationality_pattern, text, re.IGNORECASE)
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        degrees_match = re.search(degrees_pattern, text, re.IGNORECASE)
        college_name_match = re.search(college_name_pattern, text, re.IGNORECASE)
        designation_match = re.search(designation_pattern, text, re.IGNORECASE)
        last_company_match = re.search(last_company_pattern, text, re.IGNORECASE)
        job_experience_match = re.findall(job_experience_pattern, text, re.IGNORECASE)
        skills_match = re.search(skills_pattern, text, re.IGNORECASE)
        experience_match = re.findall(experience_pattern, text, re.IGNORECASE)

        # **حفظ البيانات في القاموس**
        data["name"] = name_match.group(1).strip() if name_match else ""
        data["email"] = email_match.group(0).strip() if email_match else ""
        data["phone_number"] = phone_match.group(0).strip() if phone_match else ""
        data["gender"] = gender_match.group(1).strip() if gender_match else ""
        data["nationality"] = nationality_match.group(1).strip() if nationality_match else ""
        data["linkedin"] = f"https://{linkedin_match.group(1).strip()}" if linkedin_match else ""
        data["degrees"] = degrees_match.group(1).strip() if degrees_match else ""
        data["college_name"] = college_name_match.group(0).strip() if college_name_match else ""
        data["designation"] = designation_match.group(1).strip() if designation_match else ""

        # **تحسين استخراج الشركة الأخيرة**
        if job_experience_match:
            latest_experience = job_experience_match[-1]  # آخر وظيفة مسجلة
            data["designation"] = latest_experience[0].strip()
            data["last_company"] = latest_experience[1].strip()
        elif last_company_match:
            data["last_company"] = last_company_match.group(1).strip()

        # **استخراج المهارات وتحويلها إلى قائمة**
        if skills_match:
            skills_text = skills_match.group(1)
            skills_list = re.findall(r"●\s*([\w\s&-]+)|\b([\w\s&-]+),", skills_text)
            skills_list = [s[0] or s[1] for s in skills_list if s[0] or s[1]]
            data["skills"] = list(set(skills_list))  # إزالة التكرارات

        # **حساب سنوات الخبرة**
        if experience_match:
            start_years = [int(year[0]) for year in experience_match]
            end_years = [
                int(year[1]) if year[1].isdigit() else 2025  # استبدال "Present" بسنة حالية
                for year in experience_match
            ]
            total_experience = sum(end - start for start, end in zip(start_years, end_years))
            data["total_experience_years"] = total_experience

        print(data)

        return data

    except Exception as e:
        print(f"Error extracting CV data: {e}")
        return data




from django.db.models import Q
from .models import JobDescription

def generate_report(query):
    """
    البحث عن JobDescription باستخدام العنوان أو اسم المرشح، وإرجاع تقرير يحتوي على أعلى مرشح لكل وظيفة.
    """
    job_descriptions = JobDescription.objects.filter(
        Q(title__icontains=query) | Q(rankings__cv__name__icontains=query)
    ).distinct()

    reports = []
    for job_description in job_descriptions:
        rankings = CandidateRanking.objects.filter(job_description=job_description).order_by('-overall_score')

        if rankings.exists():
            top_ranking = rankings.first()
            top_candidate_name = top_ranking.cv.name  # استرجاع اسم صاحب السيرة الذاتية
            overall_score = top_ranking.overall_score
            recommendation = top_ranking.recommendation
        else:
            top_candidate_name = None
            overall_score = None
            recommendation = None

        report = {
            "job_title": job_description.title,
            "date_created": job_description.created_at.strftime("%Y-%m-%d"),
            "top_candidate": top_candidate_name,  # إرجاع اسم المرشح بدلاً من CV_ID
            "kpi_data": {
                "overall_score": overall_score,
                "recommendation": recommendation,
            }
        }
        reports.append(report)

    return reports



import csv
from django.http import HttpResponse
from fpdf import FPDF

def export_report(report, format):
    """
    Export the report in the specified format.

    Args:
        report (dict): Report data as a dictionary.
        format (str): Desired output format ('pdf', 'csv', 'excel').

    Returns:
        str: Path to the exported file.
    """
    if format == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt=f"Report for {report['job_title']}", ln=True, align="C")
        # Add more content here...
        file_path = "/tmp/report.pdf"
        pdf.output(file_path)
        return file_path

    elif format == "csv":
        response = HttpResponse(content_type="text/csv")
        writer = csv.writer(response)
        writer.writerow(["Job Title", "Date Created", "Top Candidate"])
        writer.writerow([report["job_title"], report["date_created"], report["top_candidate"]])
        file_path = "/tmp/report.csv"
        with open(file_path, "w") as f:
            f.write(response.content.decode("utf-8"))
        return file_path

    elif format == "excel":
        # Use pandas or openpyxl for Excel export
        pass

import secrets

def generate_reset_token(user):
    """
    Generate a password reset token for the given user.

    Args:
        user (User): The user object.

    Returns:
        str: The generated token.
    """
    token = secrets.token_urlsafe(32)
    # Store the token in the database or cache with an expiration time
    return token


def validate_reset_token(token):
    """
    Validate a password reset token and retrieve the associated user.

    Args:
        token (str): The password reset token.

    Returns:
        User: The user object if valid, None otherwise.
    """
    # Query the database or cache to check if the token is valid
    # Example:
    # user = User.objects.filter(reset_token=token, token_expires__gt=timezone.now()).first()
    # return user
    return None

def apply_custom_scoring(cvs, weights):
    """
    Apply custom scoring weights to a list of CVs.

    Args:
        cvs (list): List of CV objects.
        weights (dict): Dictionary of weights (e.g., {'education': 0.3, 'experience': 0.3, 'skills': 0.4}).

    Returns:
        list: Updated list of CVs with custom scores.
    """
    for cv in cvs:
        cv.score = (
            cv.education_score * weights.get("education", 0.3) +
            cv.experience_score * weights.get("experience", 0.3) +
            cv.skills_score * weights.get("skills", 0.4)
        )
    return cvs


from .models import CV,CandidateRanking
def save_rankings(job_description, ranked_results):
    """
    Save candidate rankings to the database.

    Args:
        job_description (JobDescription): The job description object.
        ranked_results (list): Ranked results as a list of dictionaries.
    """
    for result in ranked_results:
        cv = CV.objects.get(email=result["email"])
        CandidateRanking.objects.update_or_create(
            job_description=job_description,
            cv=cv,
            defaults={
                "overall_score": result["overall_score"],
                "recommendation": result["recommendation"]
            }
        )


import PyPDF2  # للتعامل مع ملفات PDF
import docx  # للتعامل مع ملفات Word


def extract_job_description(file):
    """
    دالة لاستخراج عنوان ووصف العمل من الملف
    """
    text = ""

    # التعامل مع PDF
    if file.name.endswith('.pdf'):
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"

    # التعامل مع Word
    elif file.name.endswith('.docx'):
        doc = docx.Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"

    # تحليل النص المستخرج (قد تحتاج إلى تحسين الذكاء الاصطناعي لاستخراج العنوان والوصف)
    lines = text.split("\n")
    title = lines[0] if lines else "No Title Found"
    description = "\n".join(lines[1:]) if len(lines) > 1 else "No Description Found"

    return {'title': title.strip(), 'description': description.strip()}
